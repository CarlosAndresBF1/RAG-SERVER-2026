"""Word document parser (.doc and .docx).

Extracts text from Microsoft Word files into ParsedSection objects.
- .docx: uses python-docx to read paragraphs with heading style detection.
- .doc (legacy binary): uses antiword CLI tool as a fallback.
"""

from __future__ import annotations

import subprocess
from pathlib import Path
from typing import Optional

from odyssey_rag.ingestion.parsers.base import BaseParser, ParsedSection


class DocxParser(BaseParser):
    """Parse Word documents (.doc / .docx) into heading-delimited sections.

    For .docx files, paragraphs are grouped by heading styles (Heading 1–3).
    For legacy .doc files, text is extracted via antiword and treated as a
    single section (no heading structure available).
    """

    def parse(self, file_path: str) -> list[ParsedSection]:
        """Parse a Word document into structured sections.

        Args:
            file_path: Path to the .doc or .docx file.

        Returns:
            Ordered list of ParsedSection objects.
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix == ".doc":
            return self._parse_doc_legacy(path)

        return self._parse_docx(path)

    def supported_extensions(self) -> list[str]:
        """Return supported file extensions."""
        return [".doc", ".docx"]

    # ── .docx parsing ─────────────────────────────────────────────────────────

    def _parse_docx(self, path: Path) -> list[ParsedSection]:
        """Parse a .docx file using python-docx with heading detection."""
        import docx  # type: ignore[import-untyped]

        doc = docx.Document(str(path))

        sections: list[ParsedSection] = []
        current_section: Optional[str] = None
        current_subsection: Optional[str] = None
        current_lines: list[str] = []

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""
            text = para.text.strip()

            if not text:
                if current_lines:
                    current_lines.append("")
                continue

            # Detect heading levels from Word styles
            if style_name.startswith("heading 1") or style_name == "title":
                # Flush previous section
                self._flush_section(
                    sections, current_lines, current_section, current_subsection
                )
                current_section = text
                current_subsection = None
                current_lines = []

            elif style_name.startswith("heading 2"):
                self._flush_section(
                    sections, current_lines, current_section, current_subsection
                )
                current_subsection = text
                current_lines = []

            elif style_name.startswith("heading 3"):
                self._flush_section(
                    sections, current_lines, current_section, current_subsection
                )
                current_subsection = text
                current_lines = []

            else:
                current_lines.append(text)

        # Flush remaining content
        self._flush_section(
            sections, current_lines, current_section, current_subsection
        )

        # Also extract text from tables
        table_sections = self._extract_tables(doc, current_section)
        if table_sections:
            sections.extend(table_sections)

        return sections

    def _flush_section(
        self,
        sections: list[ParsedSection],
        lines: list[str],
        section: Optional[str],
        subsection: Optional[str],
    ) -> None:
        """Flush accumulated lines into a ParsedSection."""
        content = "\n".join(lines).strip()
        if not content:
            return
        sections.append(
            ParsedSection(
                content=content,
                section=section,
                subsection=subsection,
            )
        )

    def _extract_tables(
        self, doc: "docx.Document", default_section: Optional[str]  # type: ignore[name-defined]
    ) -> list[ParsedSection]:
        """Extract table content from the Word document as Markdown-like text."""
        import docx  # type: ignore[import-untyped]

        table_sections: list[ParsedSection] = []

        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")

            if len(rows) >= 2:
                # Insert a Markdown-style separator after the header row
                header = rows[0]
                col_count = header.count("|") - 1
                separator = "| " + " | ".join(["---"] * max(1, col_count)) + " |"
                rows.insert(1, separator)

            content = "\n".join(rows)
            if content.strip():
                table_sections.append(
                    ParsedSection(
                        content=content,
                        section=default_section,
                        subsection="Table",
                    )
                )

        return table_sections

    # ── .doc (legacy binary) parsing ──────────────────────────────────────────

    def _parse_doc_legacy(self, path: Path) -> list[ParsedSection]:
        """Parse a legacy .doc file using antiword CLI.

        Falls back to a raw byte-decoded read if antiword is not available.
        """
        text = self._extract_with_antiword(path)
        if not text:
            text = self._extract_raw_text(path)

        if not text or not text.strip():
            return []

        return [
            ParsedSection(
                content=text.strip(),
                section=path.stem,
            )
        ]

    def _extract_with_antiword(self, path: Path) -> Optional[str]:
        """Try extracting text from .doc using antiword."""
        try:
            result = subprocess.run(
                ["antiword", str(path)],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        return None

    def _extract_raw_text(self, path: Path) -> Optional[str]:
        """Last-resort extraction: read binary and decode printable text."""
        raw = path.read_bytes()
        # Try UTF-8 first, then latin-1 as fallback
        for encoding in ("utf-8", "latin-1"):
            try:
                text = raw.decode(encoding, errors="ignore")
                # Filter to only printable lines
                lines = []
                for line in text.splitlines():
                    cleaned = "".join(c for c in line if c.isprintable() or c in "\t\n")
                    cleaned = cleaned.strip()
                    if len(cleaned) > 3:
                        lines.append(cleaned)
                if lines:
                    return "\n".join(lines)
            except Exception:
                continue
        return None
