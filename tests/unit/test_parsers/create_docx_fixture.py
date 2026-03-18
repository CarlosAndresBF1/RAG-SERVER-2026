"""Generate a sample .docx fixture for testing the DocxParser."""

from pathlib import Path

import docx


def create_sample_docx(output_path: str) -> None:
    """Create a minimal .docx file with headings, paragraphs, and a table."""
    doc = docx.Document()

    doc.add_heading("Sample Document Title", level=1)
    doc.add_paragraph("This is the introduction paragraph of the sample Word document.")
    doc.add_paragraph("It contains multiple paragraphs for testing purposes.")

    doc.add_heading("Section One", level=2)
    doc.add_paragraph(
        "Section one contains detailed information about the first topic. "
        "This paragraph is intentionally longer to verify text extraction works properly "
        "across multiple sentences in a single paragraph."
    )

    doc.add_heading("Subsection A", level=3)
    doc.add_paragraph("Subsection A has specific implementation details.")

    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Section two covers a different topic entirely.")

    # Add a table
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    headers = ["Field", "Type", "Description"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "MsgId"
    table.rows[1].cells[1].text = "string"
    table.rows[1].cells[2].text = "Message identifier"
    table.rows[2].cells[0].text = "CreDtTm"
    table.rows[2].cells[1].text = "datetime"
    table.rows[2].cells[2].text = "Creation date and time"

    doc.save(output_path)


if __name__ == "__main__":
    fixtures_dir = Path(__file__).parent.parent.parent / "fixtures"
    fixtures_dir.mkdir(exist_ok=True)
    create_sample_docx(str(fixtures_dir / "sample_document.docx"))
    print("Created sample_document.docx")
